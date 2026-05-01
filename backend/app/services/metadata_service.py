"""
Metadata Extraction Service - Handles document metadata and content analysis

Extracted from api/metadata.py to follow proper service layer architecture.
"""

import os
import re
import zipfile
import xml.etree.ElementTree as ET
from datetime import datetime
from flask import current_app
from docx import Document
from docx.opc.exceptions import PackageNotFoundError

from app.core.extensions import db
from app.models import DocumentSnapshot


class MetadataService:
    """Service for extracting metadata and analyzing document content"""
    
    def __init__(self):
        pass
    
    @property
    def min_word_count(self):
        return current_app.config.get('MIN_DOCUMENT_WORDS', 50)
    
    @property
    def max_word_count(self):
        return current_app.config.get('MAX_DOCUMENT_WORDS', 15000)
    
    def extract_tracked_changes_analysis(self, file_path):
        """
        Extract REAL word additions and deletions from DOCX tracked changes.
        Analyzes w:ins (insertions) and w:del (deletions) elements to get actual contributor metrics.
        """
        try:
            with zipfile.ZipFile(file_path, 'r') as zip_file:
                if 'word/document.xml' not in zip_file.namelist():
                    return None, "No document found"
                
                doc_xml = zip_file.read('word/document.xml')
                root = ET.fromstring(doc_xml)
                
                W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
                W_AUTHOR = f'{{{W_NS}}}author'
                W_DATE = f'{{{W_NS}}}date'
                W_INS = f'{{{W_NS}}}ins'
                W_DEL = f'{{{W_NS}}}del'
                W_T = f'{{{W_NS}}}t'
                W_R = f'{{{W_NS}}}r'
                W_P = f'{{{W_NS}}}p'
                
                # Dictionary to store per-contributor metrics: author -> {words_added, words_deleted, changes_count}
                contributor_metrics = {}
                
                # Helper function to extract text from an element
                def extract_text(element):
                    """Extract all text content from an element"""
                    text_parts = []
                    for t_elem in element.findall(f'.//{W_T}'):
                        if t_elem.text:
                            text_parts.append(t_elem.text)
                    return ''.join(text_parts)
                
                # Helper function to count words in text
                def count_words(text):
                    """Count words in text"""
                    if not text:
                        return 0
                    return len(text.split())
                
                # Parse document for tracked changes
                for elem in root.iter():
                    author = None
                    change_date = None
                    words_changed = 0
                    change_type = None
                    
                    # Handle insertions (w:ins)
                    if elem.tag == W_INS:
                        author = elem.get(W_AUTHOR, '').strip()
                        change_date = elem.get(W_DATE, '').strip()
                        change_type = 'inserted'
                        
                        # Count words in inserted content
                        text_content = extract_text(elem)
                        words_changed = count_words(text_content)
                        
                    # Handle deletions (w:del)
                    elif elem.tag == W_DEL:
                        author = elem.get(W_AUTHOR, '').strip()
                        change_date = elem.get(W_DATE, '').strip()
                        change_type = 'deleted'
                        
                        # Count words in deleted content
                        text_content = extract_text(elem)
                        words_changed = count_words(text_content)
                    
                    # Record the change if we have an author
                    if author and change_type and words_changed > 0:
                        if author not in contributor_metrics:
                            contributor_metrics[author] = {
                                'name': author,
                                'words_added': 0,
                                'words_deleted': 0,
                                'insertions': 0,
                                'deletions': 0,
                                'last_change': change_date
                            }
                        
                        if change_type == 'inserted':
                            contributor_metrics[author]['words_added'] += words_changed
                            contributor_metrics[author]['insertions'] += 1
                        else:
                            contributor_metrics[author]['words_deleted'] += words_changed
                            contributor_metrics[author]['deletions'] += 1
                        
                        # Update last change date
                        if change_date:
                            if not contributor_metrics[author]['last_change'] or change_date > contributor_metrics[author]['last_change']:
                                contributor_metrics[author]['last_change'] = change_date
                
                return list(contributor_metrics.values()), None
                
        except Exception as e:
            current_app.logger.error(f"Tracked changes analysis failed: {e}")
            return None, f"Could not analyze tracked changes: {str(e)}"
    
    def extract_docx_metadata(self, file_path, external_metadata=None):
        """
        Main entry point for metadata extraction from DOCX or PDF files.
        Applies unified normalization and cleanup to all document formats.
        """
        # 1. Branch extraction based on file format
        if str(file_path).lower().endswith('.pdf'):
            metadata, parsing_error = self._extract_pdf_metadata(file_path, external_metadata)
        else:
            metadata, parsing_error = self._extract_docx_internal(file_path, external_metadata)

        # ── 2. UNIFIED NORMALIZATION (Applies to both PDF and DOCX) ──────────────────
        
        # Helper for Gmail username extraction
        def to_gmail_username(identity=None, email=None):
            candidate = (email or identity or '').strip()
            if '@' in candidate:
                return candidate.split('@', 1)[0].strip().lower()
            return (identity or '').strip()

        # [FINAL FALLBACK] Auth/Editor cleanup
        for field in ['author', 'last_editor']:
            val = str(metadata.get(field, ''))
            # Filter out generic/library names
            if not val or val.strip() == '' or 'python-docx' in val.lower() or val.lower() == 'none' or val.lower() == 'unavailable':
                metadata[field] = 'Unavailable'
            else:
                # Clean up to just the username/display name
                metadata[field] = to_gmail_username(val)

        # Cleanup contributor list: remove 'python-docx' or 'Unavailable' entries
        if metadata.get('contributors'):
            cleaned_contributors = []
            seen_names = set()
            for c in metadata['contributors']:
                name = str(c.get('name', '')).strip()
                if not name or name.lower() in ['unavailable', 'none', 'python-docx', 'unknown', '']:
                    continue
                
                # Normalize name for deduplication
                norm_name = to_gmail_username(name).lower()
                if norm_name not in seen_names:
                    cleaned_contributors.append(c)
                    seen_names.add(norm_name)
            metadata['contributors'] = cleaned_contributors[:10] # Limit to 10

        # If Author/Editor are still Unavailable, try to pick from contributors
        if metadata.get('contributors') and len(metadata['contributors']) > 0:
            if metadata['author'] == 'Unavailable':
                # Prefer Author/Owner role
                potential_author = next((c for c in metadata['contributors'] if c.get('role') in ['Author', 'Owner']), metadata['contributors'][0])
                metadata['author'] = potential_author.get('name', 'Unavailable')
            
            if metadata['last_editor'] == 'Unavailable':
                # Prefer Last Editor or Editor
                potential_editor = next((c for c in reversed(metadata['contributors']) if c.get('role') in ['Last Editor', 'Editor']), metadata['contributors'][-1])
                metadata['last_editor'] = potential_editor.get('name', 'Unavailable')

        # If author is still unavailable but we have a reliable editor, use it as fallback
        if metadata['author'] == 'Unavailable' and metadata['last_editor'] != 'Unavailable':
            metadata['author'] = metadata['last_editor']

        return metadata, parsing_error

    def _extract_docx_internal(self, file_path, external_metadata=None):
        """Internal DOCX-specific extraction logic"""
        metadata = {
            'author': 'Unavailable',
            'creation_date': None,
            'last_modified_date': None,
            'last_editor': 'Unavailable',
            'file_size': 0,
            'word_count': 0,
            'revision_count': 0,
            'editing_time_minutes': 0,
            'application': 'Unknown',
            'image_count': 0,
            'image_density_warning': False,
            'contributors': []
        }

        # [HELPER] Robust contributor deduplication and merging
        def to_gmail_username(identity=None, email=None):
            """Return Gmail username (local-part) when an email is available."""
            candidate = (email or identity or '').strip()
            if '@' in candidate:
                return candidate.split('@', 1)[0].strip().lower()
            return (identity or '').strip()

        def normalize_contributor_role(role):
            """Restrict displayed contributor roles to Author/Editor."""
            normalized = str(role or '').strip().lower()
            if normalized in ['owner', 'author']:
                return 'Author'
            if normalized in ['last editor', 'editor', 'writer', 'contributor', 'commenter', 'reader']:
                return 'Editor'
            return 'Editor'

        def add_contributor(name, role, email=None, date=None):
            if not name: return
            
            # 1. Basic cleaning
            name = str(name).strip()
            if not name or name.lower() in ['unavailable', 'none', 'python-docx', 'unknown', '']:
                return
            
            # 2. Advanced normalization (whitespace and case)
            norm_name = re.sub(r'\s+', ' ', name).strip().lower()
            norm_email = str(email).strip().lower() if email else None
            role = normalize_contributor_role(role)
            
            # 3. Check for existing record to merge or skip
            for entry in metadata['contributors']:
                # Ensure entry is a dict
                if not isinstance(entry, dict): continue
                
                existing_name = re.sub(r'\s+', ' ', str(entry.get('name', ''))).strip().lower()
                existing_email = str(entry.get('email', '')).strip().lower() if entry.get('email') else None
                
                # Match by NAME or EMAIL
                match_found = (norm_name == existing_name) or (norm_email and existing_email and norm_email == existing_email)
                
                if match_found:
                    # Enrich existing record
                    if not entry.get('email') and email: entry['email'] = email
                    
                    # Update date if provided and newer (or if existing date is missing)
                    if date:
                        if not entry.get('date'):
                            entry['date'] = date
                        else:
                            try:
                                # Compare ISO strings or datetime objects
                                if str(date) > str(entry['date']):
                                    entry['date'] = date
                            except Exception:
                                # Fallback to just overwriting if comparison fails
                                entry['date'] = date
                    
                    # Prefer Author/Editor role labels.
                    high_priority_roles = ['Author', 'Editor']
                    if role in high_priority_roles and entry.get('role') not in high_priority_roles:
                        entry['role'] = role
                    return

            # 4. Add new unique entry
            metadata['contributors'].append({
                'name': name, # Keep original casing for display
                'role': role,
                'email': email,
                'date': date
            })

        parsing_error = None
        
        try:
            # Get file size
            metadata['file_size'] = os.path.getsize(file_path)
            
            # Load document with python-docx
            doc = Document(file_path)
            
            # Extract basic properties
            core_props = doc.core_properties
            
            if core_props.author:
                metadata['author'] = core_props.author
                # Also register as a contributor for report consistency
                add_contributor(core_props.author, 'Author', date=core_props.created.isoformat() if core_props.created else None)
            
            if core_props.created:
                metadata['creation_date'] = core_props.created.isoformat()
            
            if core_props.modified:
                metadata['last_modified_date'] = core_props.modified.isoformat()
            
            if core_props.last_modified_by:
                metadata['last_editor'] = core_props.last_modified_by
                # Also register as a contributor for report consistency
                add_contributor(core_props.last_modified_by, 'Editor', date=core_props.modified.isoformat() if core_props.modified else None)
            
            if hasattr(core_props, 'revision') and core_props.revision:
                try:
                    metadata['revision_count'] = int(core_props.revision)
                except (ValueError, TypeError):
                    metadata['revision_count'] = 0
            
            # Try to extract additional metadata from XML
            try:
                with zipfile.ZipFile(file_path, 'r') as zip_file:
                    # Read app properties for more detailed metadata
                    if 'docProps/app.xml' in zip_file.namelist():
                        try:
                            app_xml = zip_file.read('docProps/app.xml')
                            app_root = ET.fromstring(app_xml)
                            all_text = {elem.tag.split('}')[-1]: elem.text for elem in app_root.iter()}
                            
                            if 'Application' in all_text:
                                metadata['application'] = all_text['Application']
                            if 'Words' in all_text:
                                try: metadata['word_count'] = int(all_text['Words'])
                                except: pass
                            if 'TotalTime' in all_text:
                                try: metadata['editing_time_minutes'] = int(all_text['TotalTime'])
                                except: pass
                        except Exception as e:
                            current_app.logger.warning(f"Error parsing app.xml: {e}")
                    
                    # Read core properties XML
                    if 'docProps/core.xml' in zip_file.namelist():
                        try:
                            core_xml = zip_file.read('docProps/core.xml')
                            core_root = ET.fromstring(core_xml)
                            for elem in core_root.iter():
                                tag = elem.tag.split('}')[-1]
                                text = elem.text
                                if not text: continue
                                
                                if tag == 'created' and not metadata['creation_date']:
                                    metadata['creation_date'] = text
                                elif tag == 'modified' and not metadata['last_modified_date']:
                                    metadata['last_modified_date'] = text
                                elif tag == 'lastModifiedBy' and metadata['last_editor'] == 'Unavailable':
                                    metadata['last_editor'] = text
                                elif tag == 'creator' and metadata['author'] == 'Unavailable':
                                    metadata['author'] = text
                                elif tag == 'contributor':
                                    add_contributor(text, 'Contributor')
                        except Exception as e:
                            current_app.logger.warning(f"Error parsing core.xml: {e}")

                    # ── Image Count Check (Anti-Screenshot Loophole) ────────────────
                    # Count files in word/media/ to see if students are hiding text in images
                    media_files = [f for f in zip_file.namelist() if f.startswith('word/media/')]
                    metadata['image_count'] = len(media_files)
                    
                    # Logic: If more than 1 image per 100 words, it's suspicious for a text-based proposal
                    word_count_for_density = metadata.get('word_count', 0)
                    if metadata['image_count'] > 5 and word_count_for_density > 0:
                        density_ratio = metadata['image_count'] / (word_count_for_density / 100)
                        if density_ratio > 1.5: # More than 1.5 images per 100 words
                            metadata['image_density_warning'] = True

                    # ── Revision-based author/last-editor extraction ──────────────────
                    # Parse tracked changes from word/document.xml and word/comments.xml.
                    # These records (w:ins, w:del, w:rPrChange, w:pPrChange, comment
                    # elements) embed the REAL person who made each change plus an ISO
                    # timestamp.  The earliest-dated person = most-likely original author;
                    # the latest-dated person = true last editor.
                    W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
                    REVISION_TAGS = {
                        f'{{{W_NS}}}ins', f'{{{W_NS}}}del',
                        f'{{{W_NS}}}rPrChange', f'{{{W_NS}}}pPrChange',
                        f'{{{W_NS}}}sectPrChange', f'{{{W_NS}}}tblPrChange',
                        f'{{{W_NS}}}trPrChange', f'{{{W_NS}}}tcPrChange',
                        f'{{{W_NS}}}comment',
                    }
                    W_AUTHOR = f'{{{W_NS}}}author'
                    W_DATE   = f'{{{W_NS}}}date'

                    # author_name -> latest ISO date string seen
                    revision_authors: dict = {}

                    for xml_part in ('word/document.xml', 'word/comments.xml'):
                        if xml_part not in zip_file.namelist():
                            continue
                        try:
                            part_root = ET.fromstring(zip_file.read(xml_part))
                            for elem in part_root.iter():
                                if elem.tag not in REVISION_TAGS:
                                    continue
                                r_author = elem.get(W_AUTHOR, '').strip()
                                r_date   = elem.get(W_DATE, '').strip()
                                if not r_author:
                                    continue
                                # Keep the latest date for each author
                                prev = revision_authors.get(r_author, '')
                                if r_date > prev:
                                    revision_authors[r_author] = r_date
                        except Exception as rev_err:
                            current_app.logger.warning(f"Revision parse error ({xml_part}): {rev_err}")

                    if revision_authors:
                        # Sort by date so we can determine earliest / latest
                        sorted_rev = sorted(revision_authors.items(), key=lambda kv: kv[1])
                        earliest_author, earliest_date = sorted_rev[0]
                        latest_author,   latest_date   = sorted_rev[-1]

                        # Revision tracking is more granular than core.xml — prefer it
                        if latest_author and latest_author.lower() not in ('unavailable', 'none'):
                            metadata['last_editor'] = latest_author

                        # Latest revision timestamp is the best signal for "Last Modified".
                        if latest_date:
                            metadata['last_modified_date'] = latest_date

                        # Use earliest revision author to fill author if still unknown
                        if metadata['author'] in ('Unavailable', '') and earliest_author:
                            if earliest_author.lower() not in ('unavailable', 'none'):
                                metadata['author'] = earliest_author

                        # Register every unique revision author as a contributor
                        for rev_name, rev_date in sorted_rev:
                            role = 'Author' if rev_name == earliest_author else 'Editor'
                            add_contributor(rev_name, role, date=rev_date or None)

            except Exception as e:
                current_app.logger.warning(f"Could not extract extended metadata: {e}")
            
        except Exception as e:
            current_app.logger.error(f"Metadata extraction failed: {e}")
            parsing_error = str(e)
            metadata['parsing_error'] = parsing_error

        # [DEDUP] Preference: External (Drive) Metadata > Internal Metadata
        if external_metadata:
            # Sync basic dates if internal ones are missing
            if not metadata['creation_date']:
                metadata['creation_date'] = external_metadata.get('createdTime')

            # Drive modifiedTime is authoritative for latest revision/edit time.
            if external_metadata.get('modifiedTime'):
                metadata['last_modified_date'] = external_metadata.get('modifiedTime')

            # [REVISION TRACKING] Store Drive revision IDs for future change detection
            if external_metadata.get('headRevisionId'):
                metadata['headRevisionId'] = external_metadata.get('headRevisionId')
            if external_metadata.get('version'):
                metadata['version'] = external_metadata.get('version')
                
            # Prefer Drive IDs for Author/Editor if internal ones are default
            if metadata['author'] == 'Unavailable':
                owners = external_metadata.get('owners', [])
                if owners:
                    owner_email = owners[0].get('emailAddress')
                    owner_name = owners[0].get('displayName')
                    metadata['author'] = to_gmail_username(owner_name, owner_email) or 'Unavailable'
            
            if metadata['last_editor'] == 'Unavailable':
                 lmu = external_metadata.get('lastModifyingUser')
                 if lmu:
                     lmu_email = lmu.get('emailAddress')
                     lmu_name = lmu.get('displayName')
                     metadata['last_editor'] = to_gmail_username(lmu_name, lmu_email) or 'Unavailable'

            # Add Drive people via the deduplicating helper
            # Owners
            for owner in external_metadata.get('owners', []):
                owner_email = owner.get('emailAddress')
                owner_name = owner.get('displayName')
                add_contributor(
                    name=to_gmail_username(owner_name, owner_email),
                    role='Owner',
                    email=owner_email,
                    date=external_metadata.get('createdTime')
                )
            
            # Last Editor (Drive)
            lmu = external_metadata.get('lastModifyingUser')
            if lmu:
                lmu_email = lmu.get('emailAddress')
                lmu_name = lmu.get('displayName')
                add_contributor(
                    name=to_gmail_username(lmu_name, lmu_email),
                    role='Last Editor',
                    email=lmu_email,
                    date=external_metadata.get('modifiedTime')
                )

            # Permissions (all collaborators)
            for perm in external_metadata.get('permissions', []):
                perm_email = perm.get('emailAddress')
                perm_name = perm.get('displayName')
                add_contributor(
                    name=to_gmail_username(perm_name, perm_email),
                    role=perm.get('role', 'contributor').capitalize(),
                    email=perm_email
                )

        # Final filesystem fallback for dates
        if not metadata['creation_date']:
            try:
                c_time = os.path.getctime(file_path)
                metadata['creation_date'] = datetime.fromtimestamp(c_time).isoformat()
            except Exception: pass
        if not metadata['last_modified_date']:
            try:
                m_time = os.path.getmtime(file_path)
                metadata['last_modified_date'] = datetime.fromtimestamp(m_time).isoformat()
            except Exception: pass

        return metadata, parsing_error

    
    def extract_document_text(self, file_path):
        """Extract full text content from DOCX or PDF file"""
        if str(file_path).lower().endswith('.pdf'):
            return self._extract_pdf_text(file_path)
            
        try:
            doc = Document(file_path)
            
            # 1. Extract from Headers and Footers
            hf_text = []
            for section in doc.sections:
                for header in [section.header, section.first_page_header, section.even_page_header]:
                    if header:
                        for p in header.paragraphs:
                            if p.text.strip(): hf_text.append(p.text.strip())
                for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
                    if footer:
                        for p in footer.paragraphs:
                            if p.text.strip(): hf_text.append(p.text.strip())

            # 2. Extract text from paragraphs (Main Body)
            paragraphs = []
            for paragraph in doc.paragraphs:
                has_page_break = False
                for run in paragraph.runs:
                    run_xml = run._element.xml
                    if 'w:type="page"' in run_xml or '<w:lastRenderedPageBreak' in run_xml:
                        has_page_break = True
                        break

                if paragraph.text.strip():
                    paragraphs.append(paragraph.text.strip())
                if has_page_break:
                    paragraphs.append('\f')
            
            # 3. Extract text from tables
            table_text = []
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            table_text.append(cell.text.strip())
            
            # 4. Extract from Shapes / Textboxes (Requires XML traversal)
            shape_text = []
            try:
                # Direct XML search for text boxes which are often missed by python-docx
                xml_content = doc._element.xml
                # Find all w:t tags (text) inside w:txbxContent (text box content)
                # This is a bit brute-force but effective for modern DOCX
                root = ET.fromstring(xml_content)
                W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
                W_T = f'{{{W_NS}}}t'
                W_TXBX = f'{{{W_NS}}}txbxContent'
                
                for txbx in root.findall(f'.//{W_TXBX}'):
                    parts = []
                    for t in txbx.findall(f'.//{W_T}'):
                        if t.text: parts.append(t.text)
                    if parts:
                        shape_text.append(' '.join(parts).strip())
            except Exception as xml_err:
                current_app.logger.warning(f"Shape text extraction failed: {xml_err}")

            # Combine all text with logical separation
            full_parts = []
            if hf_text:
                full_parts.append("--- DOCUMENT HEADERS/FOOTERS ---")
                full_parts.extend(hf_text)
            
            if paragraphs:
                full_parts.append("--- MAIN BODY ---")
                full_parts.extend(paragraphs)
            
            if table_text:
                full_parts.append("--- TABLE CONTENT ---")
                full_parts.extend(table_text)
                
            if shape_text:
                full_parts.append("--- TEXT BOXES/SHAPES ---")
                full_parts.extend(shape_text)
            
            full_text = '\n'.join(full_parts)
            return full_text, None
            
        except Exception as e:
            current_app.logger.error(f"Text extraction failed: {e}")
            return None, f"Text extraction error: {e}"
            
        except Exception as e:
            current_app.logger.error(f"Text extraction failed: {e}")
            return None, f"Text extraction error: {e}"
    
    def compute_content_statistics(self, text):
        """Compute comprehensive content statistics"""
        if not text:
            return {
                'word_count': 0,
                'character_count': 0,
                'character_count_no_spaces': 0,
                'sentence_count': 0,
                'paragraph_count': 0,
                'estimated_pages': 0,
                'average_words_per_sentence': 0,
                'average_sentence_length': 0
            }
        
        # Basic counts
        character_count = len(text)
        character_count_no_spaces = len(re.sub(r'\s+', '', text))
        
        # Word count
        words = re.findall(r"\b[\w'-]+\b", text)
        word_count = len(words)

        # Sentence count with basic abbreviation/decimal handling.
        sentence_input = re.sub(r'\s+', ' ', text).strip()
        if sentence_input:
            sentence_work = re.sub(r'(\d)\.(\d)', r'\1<prd>\2', sentence_input)
            abbreviations = [
                'Mr', 'Mrs', 'Ms', 'Dr', 'Prof', 'Sr', 'Jr', 'St',
                'etc', 'e.g', 'i.e', 'vs', 'Fig', 'No'
            ]
            for abbr in abbreviations:
                sentence_work = re.sub(
                    rf'\b{re.escape(abbr)}\.',
                    lambda m: m.group(0).replace('.', '<prd>'),
                    sentence_work,
                    flags=re.IGNORECASE
                )

            sentence_chunks = re.split(r'[.!?]+(?:\s+|$)', sentence_work)
            sentence_count = sum(
                1
                for chunk in sentence_chunks
                if re.search(r'[A-Za-z0-9]', chunk.replace('<prd>', '.'))
            )
        else:
            sentence_count = 0
        
        # Paragraph count (split by double newlines or single newlines)
        paragraphs = [p.strip() for p in text.split('\n') if p.strip()]
        paragraph_count = len(paragraphs)
        
        # Page count: explicit DOCX page breaks first, with word-based fallback.
        explicit_page_breaks = text.count('\f')
        estimated_pages = ((word_count + 249) // 250) if word_count > 0 else 0
        if explicit_page_breaks > 0:
            estimated_pages = max(explicit_page_breaks + 1, estimated_pages)
        elif word_count > 0 and estimated_pages == 0:
            estimated_pages = 1
        
        # Averages
        average_words_per_sentence = word_count / max(sentence_count, 1)
        average_sentence_length = character_count / max(sentence_count, 1)
        
        return {
            'word_count': word_count,
            'character_count': character_count,
            'character_count_no_spaces': character_count_no_spaces,
            'sentence_count': sentence_count,
            'paragraph_count': paragraph_count,
            'estimated_pages': estimated_pages,
            'page_count': estimated_pages,  # Alias for frontend compatibility
            'average_words_per_sentence': round(average_words_per_sentence, 2),
            'average_sentence_length': round(average_sentence_length, 2)
        }
    
    def validate_document_completeness(self, content_stats, text):
        """Validate document completeness according to SRS requirements"""
        # Validation warnings disabled by user request
        warnings = []
        is_complete = True
        
        # All validation checks removed to prevent warnings
        
        return is_complete, warnings
    
    def create_analysis_snapshot(self, submission, metadata, content_stats, text):
        """Create analysis snapshot for version comparison"""
        try:
            # Check for previous snapshots of the same document
            file_id = f"{submission.original_filename}_{submission.file_hash[:8]}"
            
            previous_snapshot = DocumentSnapshot.query.filter_by(
                file_id=file_id
            ).order_by(DocumentSnapshot.created_at.desc()).first()
            
            # Create new snapshot
            snapshot = DocumentSnapshot(
                file_id=file_id,
                submission_id=submission.id,
                word_count=content_stats['word_count'],
                file_hash=submission.file_hash,
                snapshot_timestamp=datetime.utcnow()
            )
            
            # Compare with previous snapshot if exists
            if previous_snapshot:
                word_count_change = content_stats['word_count'] - previous_snapshot.word_count
                if previous_snapshot.word_count > 0:
                    change_percentage = (word_count_change / previous_snapshot.word_count) * 100
                    snapshot.change_percentage = round(change_percentage, 2)
                    
                    # Flag major changes (≥50% as per SRS)
                    if abs(change_percentage) >= 50:
                        snapshot.major_changes = True
            
            db.session.add(snapshot)
            db.session.commit()
            
            return snapshot, None
            
        except Exception as e:
            current_app.logger.error(f"Failed to create snapshot: {e}")
            return None, f"Snapshot creation error: {e}"
    
    def generate_preliminary_report(self, submission, metadata, content_stats, text, is_complete, warnings):
        """Generate preliminary human-readable summary"""
        report = {
            'document_info': {
                'filename': submission.original_filename,
                'file_size_mb': round(submission.file_size / (1024 * 1024), 2),
                'submission_type': submission.submission_type,
                'submitted_at': submission.created_at.isoformat()
            },
            'metadata_summary': {
                'author': metadata.get('author', 'Unavailable'),
                'created': metadata.get('creation_date'),
                'last_modified': metadata.get('last_modified_date'),
                'last_editor': metadata.get('last_editor', 'Unavailable'),
                'revision_count': metadata.get('revision_count', 0)
            },
            'content_summary': {
                'word_count': content_stats['word_count'],
                'pages': content_stats['estimated_pages'],
                'sentences': content_stats['sentence_count'],
                'paragraphs': content_stats['paragraph_count'],
                'avg_words_per_sentence': content_stats['average_words_per_sentence']
            },
            'validation': {
                'is_complete': is_complete,
                'warnings': warnings,
                'meets_requirements': is_complete and content_stats['word_count'] >= self.min_word_count
            },
            'analysis_timestamp': datetime.utcnow().isoformat()
        }
        
        return report

    def _extract_pdf_metadata(self, file_path, external_metadata=None):
        metadata = {
            'author': 'Unavailable',
            'creation_date': None,
            'last_modified_date': None,
            'last_editor': 'Unavailable',
            'file_size': 0,
            'word_count': 0,
            'revision_count': 0,
            'editing_time_minutes': 0,
            'application': 'Unknown',
            'image_count': 0,
            'image_density_warning': False,
            'contributors': []
        }
        parsing_error = None

        try:
            metadata['file_size'] = os.path.getsize(file_path)
            from pypdf import PdfReader
            reader = PdfReader(file_path)
            meta = reader.metadata

            if meta:
                if meta.author:
                    metadata['author'] = meta.author
                    metadata['contributors'].append({'name': meta.author, 'role': 'Author'})
                if meta.creator:
                    metadata['application'] = meta.creator
                
                # Basic timestamp parsing for PDF creation/modification dates if available
                # D:YYYYMMDDHHmmSSOHH'mm' format
                def parse_pdf_date(date_str):
                    if not date_str: return None
                    date_str = str(date_str).replace('D:', '').replace("'", "")
                    try:
                        dt = datetime.strptime(date_str[:14], "%Y%m%d%H%M%S")
                        return dt.isoformat()
                    except:
                        return None

                if meta.creation_date:
                    metadata['creation_date'] = parse_pdf_date(meta.creation_date) or str(meta.creation_date)
                if meta.modification_date:
                    metadata['last_modified_date'] = parse_pdf_date(meta.modification_date) or str(meta.modification_date)
            
            # Count images in PDF
            img_count = 0
            for page in reader.pages:
                if '/XObject' in page['/Resources']:
                    xObject = page['/Resources']['/XObject'].get_object()
                    for obj in xObject:
                        if xObject[obj]['/Subtype'] == '/Image':
                            img_count += 1
            metadata['image_count'] = img_count

        except Exception as e:
            current_app.logger.error(f"PDF metadata extraction failed: {e}")
            parsing_error = str(e)
            metadata['parsing_error'] = parsing_error

        # Deduplicate and merge external metadata
        if external_metadata:
            if not metadata['creation_date']:
                metadata['creation_date'] = external_metadata.get('createdTime')
            if external_metadata.get('modifiedTime'):
                metadata['last_modified_date'] = external_metadata.get('modifiedTime')
            
            if metadata['author'] == 'Unavailable':
                owners = external_metadata.get('owners', [])
                if owners:
                    metadata['author'] = owners[0].get('displayName') or owners[0].get('emailAddress') or 'Unavailable'
            
            if metadata['last_editor'] == 'Unavailable':
                lmu = external_metadata.get('lastModifyingUser')
                if lmu:
                    metadata['last_editor'] = lmu.get('displayName') or lmu.get('emailAddress') or 'Unavailable'

        if not metadata['creation_date']:
            try: metadata['creation_date'] = datetime.fromtimestamp(os.path.getctime(file_path)).isoformat()
            except: pass
        if not metadata['last_modified_date']:
            try: metadata['last_modified_date'] = datetime.fromtimestamp(os.path.getmtime(file_path)).isoformat()
            except: pass

        return metadata, parsing_error

    def _extract_pdf_text(self, file_path):
        try:
            from pypdf import PdfReader
            reader = PdfReader(file_path)
            text_parts = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
            return "\n\f\n".join(text_parts), None
        except Exception as e:
            current_app.logger.error(f"PDF text extraction failed: {e}")
            return None, f"PDF text extraction error: {e}"
